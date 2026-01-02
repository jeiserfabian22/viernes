# software/utils/firma_digital.py

import hashlib
import base64
import io
import os
from datetime import datetime, timedelta
from PIL import Image, ImageDraw, ImageFont
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from cryptography.hazmat.primitives import hashes, serialization
from cryptography.hazmat.primitives.asymmetric import rsa, padding
from cryptography.hazmat.backends import default_backend
from cryptography.exceptions import InvalidSignature


class FirmaDigitalManager:
    """Gestor de firma digital criptográfica"""
    
    @staticmethod
    def generar_certificado_usuario(usuario):
        """
        Genera un certificado digital para un usuario
        Retorna: (clave_privada_pem, clave_publica_pem, huella_digital)
        """
        # Generar par de claves RSA 2048 bits
        private_key = rsa.generate_private_key(
            public_exponent=65537,
            key_size=2048,
            backend=default_backend()
        )
        
        # Serializar clave privada (sin encriptar para simplificar)
        private_pem = private_key.private_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PrivateFormat.PKCS8,
            encryption_algorithm=serialization.NoEncryption()
        )
        
        # Serializar clave pública
        public_key = private_key.public_key()
        public_pem = public_key.public_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PublicFormat.SubjectPublicKeyInfo
        )
        
        # Calcular huella digital (SHA-256 de la clave pública)
        huella = hashlib.sha256(public_pem).hexdigest()
        
        print(f"✅ Certificado generado para: {usuario.nombrecompleto}")
        print(f"   Huella digital: {huella[:16]}...")
        
        return (
            private_pem.decode('utf-8'),
            public_pem.decode('utf-8'),
            huella
        )
    
    @staticmethod
    def calcular_hash_archivo(archivo):
        """
        Calcula SHA-256 de un archivo
        """
        sha256 = hashlib.sha256()
        archivo.seek(0)
        
        for chunk in iter(lambda: archivo.read(8192), b''):
            sha256.update(chunk)
        
        archivo.seek(0)
        hash_hex = sha256.hexdigest()
        
        print(f"📄 Hash calculado: {hash_hex[:16]}...")
        return hash_hex
    
    @staticmethod
    def firmar_hash(hash_documento, clave_privada_pem):
        """
        Firma un hash con la clave privada
        Retorna: firma en base64
        """
        # Cargar clave privada
        private_key = serialization.load_pem_private_key(
            clave_privada_pem.encode(),
            password=None,
            backend=default_backend()
        )
        
        # Firmar el hash
        signature = private_key.sign(
            hash_documento.encode('utf-8'),
            padding.PSS(
                mgf=padding.MGF1(hashes.SHA256()),
                salt_length=padding.PSS.MAX_LENGTH
            ),
            hashes.SHA256()
        )
        
        # Convertir a base64
        firma_b64 = base64.b64encode(signature).decode('utf-8')
        
        print(f"✍️ Firma digital generada: {firma_b64[:32]}...")
        return firma_b64
    
    @staticmethod
    def verificar_firma(hash_documento, firma_b64, clave_publica_pem):
        """
        Verifica la firma digital
        Retorna: True si es válida, False si fue modificado
        """
        try:
            # Cargar clave pública
            public_key = serialization.load_pem_public_key(
                clave_publica_pem.encode(),
                backend=default_backend()
            )
            
            # Decodificar firma
            signature = base64.b64decode(firma_b64)
            
            # Verificar firma
            public_key.verify(
                signature,
                hash_documento.encode('utf-8'),
                padding.PSS(
                    mgf=padding.MGF1(hashes.SHA256()),
                    salt_length=padding.PSS.MAX_LENGTH
                ),
                hashes.SHA256()
            )
            
            print("✅ Firma VÁLIDA - Documento íntegro")
            return True
            
        except InvalidSignature:
            print("❌ Firma INVÁLIDA - Documento modificado")
            return False
        except Exception as e:
            print(f"❌ Error en verificación: {e}")
            return False
    
    @staticmethod
    def agregar_firma_visual_pdf(archivo_pdf, firma_canvas_b64, metadata):
        """
        Agrega firma visual a un PDF en la esquina SUPERIOR DERECHA de la PRIMERA página
        """
        from io import BytesIO
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import letter
        
        # Leer PDF original
        pdf_reader = PdfReader(archivo_pdf)
        pdf_writer = PdfWriter()
        
        # Decodificar imagen de firma
        firma_img_data = base64.b64decode(firma_canvas_b64.split(',')[1])
        firma_img = Image.open(BytesIO(firma_img_data))
        
        # Crear overlay con firma
        packet = BytesIO()
        can = rl_canvas.Canvas(packet, pagesize=letter)
        
        # ⭐ Obtener dimensiones de la PRIMERA página
        page = pdf_reader.pages[0]  # PRIMERA PÁGINA
        page_width = float(page.mediabox.width)
        page_height = float(page.mediabox.height)
        
        # ⭐ POSICIÓN: ESQUINA SUPERIOR DERECHA
        firma_width = 120   # Ancho reducido
        firma_height = 100  # Alto reducido
        margen_derecho = 50
        margen_superior = 50
        
        x_pos = page_width - firma_width - margen_derecho
        y_pos = page_height - firma_height - margen_superior  # ⭐ Desde arriba
        
        # Guardar imagen temporalmente
        temp_img = BytesIO()
        firma_img.save(temp_img, format='PNG')
        temp_img.seek(0)
        
        # Dibujar firma en canvas
        can.drawImage(
            ImageReader(temp_img),
            x_pos, y_pos,
            width=firma_width,
            height=firma_height,
            mask='auto',
            preserveAspectRatio=True
        )
        
        # ⭐ Agregar metadatos DEBAJO de la firma (más pequeños)
        can.setFont("Helvetica", 6)  # Fuente más pequeña
        can.setFillColorRGB(0.4, 0.4, 0.4)  # Gris más claro
        
        text_y = y_pos - 8  # Justo debajo de la firma
        can.drawString(x_pos, text_y, f"Firmado por: {metadata['usuario']}")
        can.drawString(x_pos, text_y - 10, f"Fecha: {metadata['fecha']}")
        
        # Código de verificación más pequeño
        can.setFont("Helvetica", 5)
        can.drawString(x_pos, text_y - 20, f"Cert: {metadata['huella'][:12]}...")
        
        can.save()
        
        # Crear página de overlay
        packet.seek(0)
        overlay_pdf = PdfReader(packet)
        overlay_page = overlay_pdf.pages[0]
        
        # ⭐ Fusionar SOLO en la PRIMERA página
        for i, page in enumerate(pdf_reader.pages):
            if i == 0:  # PRIMERA PÁGINA
                page.merge_page(overlay_page)
            pdf_writer.add_page(page)
        
        # Agregar metadatos al PDF
        pdf_writer.add_metadata({
            '/Producer': 'MotoVentas - Sistema de Firma Digital',
            '/Creator': metadata['usuario'],
            '/CreationDate': metadata['fecha'],
            '/Title': f"Documento firmado digitalmente",
            '/Subject': 'Firma Digital Criptográfica'
        })
        
        # Escribir PDF resultante
        output = BytesIO()
        pdf_writer.write(output)
        output.seek(0)
        
        print(f"✅ Firma visual agregada en esquina superior derecha - Primera página")
        return output
    
    @staticmethod
    def agregar_firma_visual_docx(archivo_docx, firma_canvas_b64, metadata):
        """
        Agrega firma visual en la PRIMERA PÁGINA del documento Word
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Leer documento original
        archivo_docx.seek(0)
        doc = Document(archivo_docx)
        
        # Decodificar imagen de firma
        firma_img_data = base64.b64decode(firma_canvas_b64.split(',')[1])
        firma_img_bytes = io.BytesIO(firma_img_data)
        
        # ⭐ INSERTAR TABLA AL INICIO DEL DOCUMENTO (primera página)
        # Crear tabla ANTES del primer párrafo
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        
        # ⭐ MOVER la tabla al INICIO del documento
        # Obtener el elemento de la tabla que acabamos de crear
        table_element = table._element
        
        # Obtener el cuerpo del documento
        body = doc._body._element
        
        # Insertar la tabla al PRINCIPIO (posición 0)
        body.insert(0, table_element)
        
        # Configurar anchos de columnas
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]
        
        left_cell.width = Inches(4.0)
        right_cell.width = Inches(2.5)
        
        # Celda izquierda vacía
        left_cell.text = ""
        
        # Celda derecha con firma
        paragraph_firma = right_cell.paragraphs[0]
        paragraph_firma.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Agregar imagen de firma
        run = paragraph_firma.add_run()
        run.add_picture(firma_img_bytes, width=Inches(1.3))
        
        # Agregar metadatos
        paragraph_meta = right_cell.add_paragraph()
        paragraph_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run_meta = paragraph_meta.add_run(
            f"Firmado digitalmente\n"
            f"{metadata['usuario']}\n"
            f"{metadata['fecha']}\n"
            f"Cert: {metadata['huella'][:16]}..."
        )
        run_meta.font.size = Pt(7)
        run_meta.font.color.rgb = RGBColor(100, 100, 100)
        
        # ⭐ Quitar bordes de la tabla
        def remove_table_borders(table):
            """Función para quitar bordes de tabla"""
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Crear elemento de bordes
            tblBorders = OxmlElement('w:tblBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tblBorders.append(border)
            
            tblPr.append(tblBorders)
        
        # Aplicar función para quitar bordes
        try:
            remove_table_borders(table)
        except Exception as e:
            print(f"⚠️ No se pudieron quitar los bordes de la tabla: {e}")
        
        # ⭐ Agregar un párrafo vacío después de la firma para separar del contenido
        separator = doc.add_paragraph()
        # Mover el separador al inicio también (después de la tabla)
        separator_element = separator._element
        body.insert(1, separator_element)
        
        # Guardar documento
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        print(f"✅ Firma visual agregada al INICIO del documento Word")
        return output

    
    @staticmethod
    def generar_codigo_verificacion(iddocumento):
        """
        Genera un código de verificación único
        """
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        codigo = f"MV-{iddocumento:06d}-{timestamp}"
        return codigo